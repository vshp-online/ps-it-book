#!/usr/bin/env python3
"""Создаёт A4-шаблон DOCX для редакторского экспорта книги.

Запускать из корня репозитория после получения стандартного reference.docx:

    quarto pandoc --print-default-data-file reference.docx > tmp/reference.docx
    python3 scripts/build-editorial-reference.py \
        tmp/reference.docx book/templates/editorial-reference.docx
"""

from pathlib import Path
import sys

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


BODY_FONT = "Liberation Serif"
HEADING_FONT = "Liberation Sans"
CODE_FONT = "Noto Sans Mono"
ACCENT = RGBColor(31, 90, 112)
DARK_GRAY = RGBColor(70, 70, 70)


def set_run_fonts(run, name: str) -> None:
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)


def set_style_font(style, name: str, size: float, *, bold=None, italic=None, color=None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    if color is not None:
        style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)


def set_outline_level(style, level: int) -> None:
    ppr = style.element.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(level))


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_fonts(run, BODY_FONT)
    run.font.size = Pt(9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    set_style_font(normal, BODY_FONT, 11.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    set_style_font(title, HEADING_FONT, 24, bold=True, color=ACCENT)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, HEADING_FONT, 13, color=DARK_GRAY)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)

    if "Author" in styles:
        author = styles["Author"]
        set_style_font(author, BODY_FONT, 11)
        author.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading_specs = {
        "Heading 1": (17, 16, 6, 1),
        "Heading 2": (13.5, 12, 4, 2),
        "Heading 3": (11.5, 9, 3, 3),
        "Heading 4": (11, 7, 2, 4),
    }
    for name, (size, before, after, outline) in heading_specs.items():
        if name not in styles:
            continue
        style = styles[name]
        set_style_font(style, HEADING_FONT, size, bold=True, color=ACCENT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        set_outline_level(style, outline)
    styles["Heading 1"].paragraph_format.page_break_before = True

    if "Part Title" in styles:
        part = styles["Part Title"]
    else:
        part = styles.add_style("Part Title", WD_STYLE_TYPE.PARAGRAPH)
    part.base_style = styles["Title"]
    set_style_font(part, HEADING_FONT, 21, bold=True, color=ACCENT)
    part.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    part.paragraph_format.space_before = Pt(0)
    part.paragraph_format.space_after = Pt(18)
    part.paragraph_format.page_break_before = True
    part.paragraph_format.keep_with_next = True
    part.paragraph_format.keep_together = True
    set_outline_level(part, 0)

    for name in ("Image Caption", "Table Caption", "Caption"):
        if name in styles:
            style = styles[name]
            set_style_font(style, BODY_FONT, 10.5, italic=True)
            style.paragraph_format.space_before = Pt(3)
            style.paragraph_format.space_after = Pt(8)
            style.paragraph_format.keep_together = True

    for name in ("Source Code", "Verbatim Char"):
        if name in styles:
            set_style_font(styles[name], CODE_FONT, 8.5)

    if "Footnote Text" in styles:
        set_style_font(styles["Footnote Text"], BODY_FONT, 9)

    if "TOC Heading" in styles:
        toc_heading = styles["TOC Heading"]
        set_style_font(toc_heading, HEADING_FONT, 17, bold=True, color=ACCENT)
        toc_heading.paragraph_format.page_break_before = True
        toc_heading.paragraph_format.space_after = Pt(8)

    for level in range(1, 5):
        name = f"TOC {level}"
        if name not in styles:
            continue
        style = styles[name]
        set_style_font(style, BODY_FONT, max(9.5, 11.5 - 0.5 * (level - 1)))
        style.paragraph_format.space_after = Pt(2)


def configure_page(document: Document) -> None:
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(22)
        section.bottom_margin = Mm(22)
        section.left_margin = Mm(25)
        section.right_margin = Mm(20)
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)
        footer = section.footer
        paragraph = footer.paragraphs[0]
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)
        add_page_field(paragraph)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Использование: build-editorial-reference.py INPUT OUTPUT")
    source = Path(sys.argv[1])
    target = Path(sys.argv[2])
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document(source)
    configure_page(document)
    configure_styles(document)
    document.save(target)
    print(target)


if __name__ == "__main__":
    main()
